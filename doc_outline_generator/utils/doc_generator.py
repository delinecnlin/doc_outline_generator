import os
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
from typing import Dict

class DocGenerator:
    def __init__(self, outline: str, project_info: Dict[str, str]):
        self.outline = outline
        self.project_info = project_info

    def generate_doc(self) -> str:
        doc = Document()
        self._add_outline(doc)
        self._add_project_info(doc)
        self._add_diagram(doc)

        doc_path = 'generated_doc.docx'
        doc.save(doc_path)

        return doc_path

    def _add_outline(self, doc: Document):
        doc.add_heading('Project Outline', 0)

        for idx, item in enumerate(self.outline.split('\n'), 1):
            doc.add_heading(f'Part {idx}', level=1)
            doc.add_paragraph(item)

        doc.add_page_break()

    def _add_project_info(self, doc: Document):
        doc.add_heading('Project Information', 0)
        for key, value in self.project_info.items():
            doc.add_heading(key, level=1)
            doc.add_paragraph(value)

        doc.add_page_break()

    def _add_diagram(self, doc: Document):
        doc.add_heading('System Design Diagram', 0)
        fig = self._generate_diagram()
        fig.savefig('diagram.png', dpi=300)
        doc.add_picture('diagram.png', width=Inches(6.0))

    def _generate_diagram(self):
        fig, ax = plt.subplots()
        sns.heatmap([[0, 1], [1, 0]], ax=ax, cmap='Blues', annot=True)
        return fig
